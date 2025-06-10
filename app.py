from flask import Flask, request, render_template, send_file
from PIL import Image
from PyPDF2 import PdfReader
from docx import Document
import os
import requests
import io
import traceback

app = Flask(__name__, template_folder='.')
UPLOAD_FOLDER = '/tmp/uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)

@app.route('/')
def index():
    try:
        return render_template('index.html')
    except Exception as e:
        return f"Error rendering template: {str(e)}", 500

@app.route('/convert', methods=['POST'])
def convert():
    try:
        if 'file' not in request.files:
            return "No file uploaded", 400
        file = request.files['file']
        conversion_type = request.form.get('conversion_type', '')
        if file.filename == '':
            return "No file selected", 400
        if not conversion_type:
            return "No conversion type selected", 400

        filename = file.filename.replace(' ', '_').replace('/', '_')
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(input_path)
        output_filename = 'converted_' + filename
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)

        if conversion_type in ['jpg_to_png', 'png_to_jpg', 'webp_to_png', 'png_to_webp', 'gif_to_png', 'tiff_to_jpg']:
            img = Image.open(input_path)
            img = img.convert('RGB') if conversion_type != 'png_to_webp' else img
            if conversion_type == 'jpg_to_png':
                output_path = output_path.replace('.jpg', '.png').replace('.jpeg', '.png')
                img.save(output_path, 'PNG')
            elif conversion_type == 'png_to_jpg':
                output_path = output_path.replace('.png', '.jpg')
                img.save(output_path, 'JPEG', quality=95)
            elif conversion_type == 'webp_to_png':
                output_path = output_path.replace('.webp', '.png')
                img.save(output_path, 'PNG')
            elif conversion_type == 'png_to_webp':
                output_path = output_path.replace('.png', '.webp')
                img.save(output_path, 'WEBP')
            elif conversion_type == 'gif_to_png':
                output_path = output_path.replace('.gif', '.png')
                img.save(output_path, 'PNG')
            elif conversion_type == 'tiff_to_jpg':
                output_path = output_path.replace('.tiff', '.jpg').replace('.tif', '.jpg')
                img.save(output_path, 'JPEG', quality=95)
        elif conversion_type == 'pdf_to_docx':
            pdf = PdfReader(input_path)
            doc = Document()
            for page in pdf.pages:
                text = page.extract_text() or ''
                doc.add_paragraph(text)
            output_path = output_path.replace('.pdf', '.docx')
            doc.save(output_path)
        elif conversion_type == 'docx_to_pdf':
            doc = Document(input_path)
            html_content = "<html><body>"
            for para in doc.paragraphs:
                html_content += f"<p>{para.text}</p>"
            html_content += "</body></html>"
            output_path = output_path.replace('.docx', '.html')
            with open(output_path, 'w') as f:
                f.write(html_content)
        elif conversion_type == 'remove_bg':
            api_key = os.getenv('REMOVE_BG_API_KEY', '')
            if not api_key:
                return "Remove.bg API key not configured", 500
            url = 'https://api.remove.bg/v1.0/removebg'
            files = {'image_file': open(input_path, 'rb')}
            headers = {'X-Api-Key': api_key}
            response = requests.post(url, files=files, headers=headers, timeout=10)
            if response.status_code == 200:
                output_path = output_path.rsplit('.', 1)[0] + '_nobg.png'
                with open(output_path, 'wb') as out:
                    out.write(response.content)
            else:
                return f"Background removal failed: {response.text}", 500
        else:
            return "Invalid conversion type", 400

        if not os.path.exists(output_path):
            return "Output file not created", 500
        return send_file(output_path, as_attachment=True, download_name=output_filename)
    except Exception as e:
        traceback.print_exc()
        return f"Error during conversion: {str(e)}", 500
    finally:
        if os.path.exists(input_path):
            try:
                os.remove(input_path)
            except:
                pass
        if os.path.exists(output_path):
            try:
                os.remove(output_path)
            except:
                pass

if __name__ == '__main__':
    app.run(debug=True)
